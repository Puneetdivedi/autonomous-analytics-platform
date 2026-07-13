"""Report document generators (PDF / DOCX / XLSX).

Each generator accepts a title plus a list of ``sections`` and returns the
absolute path of the written file. A section is a dict shaped like::

    {
        "heading": "Executive Summary",
        "body": "Free-text paragraph(s).",
        "bullets": ["point one", "point two"],
        "table": {"columns": ["a", "b"], "rows": [{"a": 1, "b": 2}]},
    }

All keys are optional. Charts are passed separately as base64-encoded PNG
strings and embedded where the format supports images (PDF, DOCX).
"""

from __future__ import annotations

import base64
import io
from pathlib import Path
from typing import Any

from app.core.exceptions import DataSourceError
from app.core.logging import get_logger
from app.tools.file_writer import resolve_path

logger = get_logger(__name__)


def _ensure_parent(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)


def _section_table(section: dict[str, Any]) -> tuple[list[str], list[list[str]]] | None:
    table = section.get("table")
    if not isinstance(table, dict):
        return None
    columns = [str(c) for c in table.get("columns", [])]
    rows_in = table.get("rows", [])
    if not columns and rows_in:
        columns = [str(k) for k in rows_in[0].keys()]
    rows = [[str(row.get(col, "")) for col in columns] for row in rows_in]
    return columns, rows


def generate_pdf(
    title: str,
    sections: list[dict[str, Any]],
    charts_b64: list[str] | None,
    out_path: str | Path,
) -> str:
    """Render a PDF report with reportlab. Returns the absolute file path."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        Image,
        ListFlowable,
        ListItem,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    target = resolve_path(out_path, "report")
    try:
        _ensure_parent(target)
        styles = getSampleStyleSheet()
        doc = SimpleDocTemplate(str(target), pagesize=letter)
        story: list[Any] = [Paragraph(title, styles["Title"]), Spacer(1, 0.25 * inch)]

        for section in sections or []:
            if section.get("heading"):
                story.append(Paragraph(str(section["heading"]), styles["Heading2"]))
            if section.get("body"):
                story.append(Paragraph(str(section["body"]), styles["BodyText"]))
            bullets = section.get("bullets") or []
            if bullets:
                story.append(
                    ListFlowable(
                        [ListItem(Paragraph(str(b), styles["BodyText"])) for b in bullets],
                        bulletType="bullet",
                    )
                )
            table = _section_table(section)
            if table:
                columns, rows = table
                data = [columns, *rows] if columns else rows
                if data:
                    tbl = Table(data)
                    tbl.setStyle(
                        TableStyle(
                            [
                                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2d3748")),
                                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                                ("FONTSIZE", (0, 0), (-1, -1), 8),
                            ]
                        )
                    )
                    story.append(tbl)
            story.append(Spacer(1, 0.2 * inch))

        for encoded in charts_b64 or []:
            try:
                image_bytes = base64.b64decode(encoded)
                story.append(Image(io.BytesIO(image_bytes), width=6 * inch, height=3.75 * inch))
                story.append(Spacer(1, 0.2 * inch))
            except Exception as exc:  # noqa: BLE001
                logger.warning("pdf_chart_embed_failed", error=str(exc))

        doc.build(story)
        logger.info("pdf_generated", path=str(target))
        return str(target)
    except DataSourceError:
        raise
    except Exception as exc:  # noqa: BLE001
        logger.error("pdf_generation_failed", error=str(exc))
        raise DataSourceError(f"Failed to generate PDF: {exc}") from exc


def generate_docx(
    title: str,
    sections: list[dict[str, Any]],
    charts_b64: list[str] | None,
    out_path: str | Path,
) -> str:
    """Render a DOCX report with python-docx. Returns the absolute file path."""
    from docx import Document
    from docx.shared import Inches

    target = resolve_path(out_path, "report")
    try:
        _ensure_parent(target)
        document = Document()
        document.add_heading(title, level=0)

        for section in sections or []:
            if section.get("heading"):
                document.add_heading(str(section["heading"]), level=1)
            if section.get("body"):
                document.add_paragraph(str(section["body"]))
            for bullet in section.get("bullets") or []:
                document.add_paragraph(str(bullet), style="List Bullet")
            table = _section_table(section)
            if table:
                columns, rows = table
                if columns:
                    doc_table = document.add_table(rows=1, cols=len(columns))
                    doc_table.style = "Light Grid Accent 1"
                    header = doc_table.rows[0].cells
                    for i, col in enumerate(columns):
                        header[i].text = col
                    for row in rows:
                        cells = doc_table.add_row().cells
                        for i, val in enumerate(row):
                            cells[i].text = val

        for encoded in charts_b64 or []:
            try:
                image_bytes = base64.b64decode(encoded)
                document.add_picture(io.BytesIO(image_bytes), width=Inches(6))
            except Exception as exc:  # noqa: BLE001
                logger.warning("docx_chart_embed_failed", error=str(exc))

        document.save(str(target))
        logger.info("docx_generated", path=str(target))
        return str(target)
    except DataSourceError:
        raise
    except Exception as exc:  # noqa: BLE001
        logger.error("docx_generation_failed", error=str(exc))
        raise DataSourceError(f"Failed to generate DOCX: {exc}") from exc


def generate_xlsx(
    title: str,
    sections: list[dict[str, Any]],
    out_path: str | Path,
) -> str:
    """Write a workbook: a summary sheet plus one sheet per tabular section."""
    from openpyxl import Workbook
    from openpyxl.styles import Font

    target = resolve_path(out_path, "report")
    try:
        _ensure_parent(target)
        wb = Workbook()
        summary = wb.active
        summary.title = "Summary"
        summary["A1"] = title
        summary["A1"].font = Font(bold=True, size=14)

        row_cursor = 3
        table_index = 0
        for section in sections or []:
            heading = str(section.get("heading", "")) or None
            if heading:
                summary.cell(row=row_cursor, column=1, value=heading).font = Font(bold=True)
                row_cursor += 1
            if section.get("body"):
                summary.cell(row=row_cursor, column=1, value=str(section["body"]))
                row_cursor += 1
            for bullet in section.get("bullets") or []:
                summary.cell(row=row_cursor, column=1, value=f"- {bullet}")
                row_cursor += 1
            row_cursor += 1

            table = _section_table(section)
            if table and table[0]:
                table_index += 1
                columns, rows = table
                sheet_name = (heading or f"Data {table_index}")[:28] or f"Data {table_index}"
                # Ensure uniqueness of sheet names.
                base_name = sheet_name
                suffix = 1
                while sheet_name in wb.sheetnames:
                    suffix += 1
                    sheet_name = f"{base_name[:26]}_{suffix}"
                ws = wb.create_sheet(title=sheet_name)
                ws.append(columns)
                for cell in ws[1]:
                    cell.font = Font(bold=True)
                for row in rows:
                    ws.append(row)

        wb.save(str(target))
        logger.info("xlsx_generated", path=str(target))
        return str(target)
    except DataSourceError:
        raise
    except Exception as exc:  # noqa: BLE001
        logger.error("xlsx_generation_failed", error=str(exc))
        raise DataSourceError(f"Failed to generate XLSX: {exc}") from exc


__all__ = ["generate_pdf", "generate_docx", "generate_xlsx"]
